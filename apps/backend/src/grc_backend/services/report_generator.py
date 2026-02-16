"""Report generation service for GRC documents.

Generates various types of reports from interview transcripts:
- Business Process Documents (業務記述書)
- Risk Control Matrix (RCM)
- Audit Workpapers (監査調書)
- Interview Summaries
"""

import logging
from dataclasses import dataclass, field
from datetime import datetime
from enum import StrEnum
from typing import Any
from uuid import UUID

logger = logging.getLogger(__name__)


class ReportType(StrEnum):
    """Types of reports that can be generated."""

    SUMMARY = "summary"  # Interview summary
    PROCESS_DOC = "process_doc"  # Business process document
    RCM = "rcm"  # Risk Control Matrix
    AUDIT_WORKPAPER = "audit_workpaper"  # Audit workpaper


class ExportFormat(StrEnum):
    """Supported export formats."""

    JSON = "json"
    MARKDOWN = "markdown"
    WORD = "word"
    EXCEL = "excel"
    PDF = "pdf"


@dataclass
class ProcessStep:
    """A step in a business process."""

    step_number: int
    description: str
    responsible_party: str | None = None
    system_used: str | None = None
    inputs: list[str] = field(default_factory=list)
    outputs: list[str] = field(default_factory=list)
    controls: list[str] = field(default_factory=list)


@dataclass
class RiskControlItem:
    """An item in the Risk Control Matrix."""

    risk_id: str
    risk_description: str
    risk_category: str
    likelihood: str  # High, Medium, Low
    impact: str  # High, Medium, Low
    control_id: str
    control_description: str
    control_type: str  # Preventive, Detective, Corrective
    control_frequency: str  # Daily, Weekly, Monthly, etc.
    control_owner: str | None = None
    residual_risk: str | None = None


@dataclass
class AuditFinding:
    """An audit finding."""

    finding_id: str
    finding_type: str  # Observation, Finding, Exception
    description: str
    criteria: str
    condition: str
    cause: str | None = None
    effect: str | None = None
    recommendation: str | None = None
    management_response: str | None = None


@dataclass
class GeneratedReport:
    """Result of report generation."""

    report_type: ReportType
    title: str
    generated_at: datetime
    content: dict[str, Any]
    metadata: dict[str, Any] = field(default_factory=dict)


class ReportGeneratorService:
    """Service for generating GRC reports from interview data.

    Uses AI models to analyze interview transcripts and generate
    structured documents for GRC advisory work.
    """

    # Prompts for different report types
    PROMPTS = {
        ReportType.SUMMARY: """
あなたはGRCコンサルタントです。以下のインタビュー記録を分析し、
構造化された要約レポートを作成してください。

## 出力形式 (JSON)
{{
    "title": "インタビュー要約",
    "interviewee_role": "被インタビュー者の役職・役割",
    "date": "インタビュー日",
    "duration_minutes": インタビュー時間（分）,
    "key_topics": ["主要トピック1", "主要トピック2"],
    "summary": "全体要約（3-5文）",
    "key_findings": [
        {{
            "topic": "トピック",
            "finding": "発見事項",
            "significance": "重要度（高/中/低）"
        }}
    ],
    "follow_up_items": ["フォローアップ項目1", "フォローアップ項目2"],
    "quotes": ["重要な発言の引用"]
}}

## インタビュー記録
{transcript}
""",
        ReportType.PROCESS_DOC: """
あなたはGRCコンサルタントです。以下のインタビュー記録を分析し、
業務記述書（ナラティブ形式）を作成してください。

## 出力形式 (JSON)
{{
    "title": "業務記述書",
    "process_name": "業務プロセス名",
    "process_owner": "プロセスオーナー",
    "department": "担当部門",
    "objective": "業務目的",
    "scope": "対象範囲",
    "frequency": "実施頻度",
    "narrative": "業務の流れを説明するナラティブ（段落形式）",
    "process_steps": [
        {{
            "step_number": 1,
            "description": "ステップの説明",
            "responsible_party": "担当者/部門",
            "system_used": "使用システム",
            "inputs": ["入力1", "入力2"],
            "outputs": ["出力1", "出力2"],
            "controls": ["関連する統制"]
        }}
    ],
    "key_controls": ["主要な統制1", "主要な統制2"],
    "systems_involved": ["関連システム1", "関連システム2"],
    "risks": ["識別されたリスク1", "識別されたリスク2"],
    "improvement_opportunities": ["改善機会1", "改善機会2"]
}}

## インタビュー記録
{transcript}
""",
        ReportType.RCM: """
あなたはGRCコンサルタントです。以下のインタビュー記録を分析し、
リスクコントロールマトリックス（RCM）を作成してください。

## 出力形式 (JSON)
{{
    "title": "リスクコントロールマトリックス",
    "process_name": "対象業務プロセス",
    "assessment_date": "評価日",
    "prepared_by": "作成者",
    "items": [
        {{
            "risk_id": "R001",
            "risk_description": "リスクの説明",
            "risk_category": "カテゴリ（財務報告/業務/コンプライアンス）",
            "likelihood": "発生可能性（高/中/低）",
            "impact": "影響度（高/中/低）",
            "inherent_risk": "固有リスク評価",
            "control_id": "C001",
            "control_description": "統制の説明",
            "control_type": "統制タイプ（予防的/発見的/是正的）",
            "control_frequency": "実施頻度",
            "control_owner": "統制責任者",
            "control_effectiveness": "統制の有効性評価",
            "residual_risk": "残余リスク評価",
            "gap_identified": "識別されたギャップ",
            "remediation_action": "改善アクション"
        }}
    ],
    "summary": {{
        "total_risks": リスク総数,
        "high_risks": 高リスク数,
        "medium_risks": 中リスク数,
        "low_risks": 低リスク数,
        "control_gaps": ギャップ数,
        "key_observations": ["主要な所見"]
    }}
}}

## インタビュー記録
{transcript}
""",
        ReportType.AUDIT_WORKPAPER: """
あなたは内部監査人です。以下のインタビュー記録を分析し、
監査調書を作成してください。

## 出力形式 (JSON)
{{
    "title": "監査調書",
    "audit_project": "監査プロジェクト名",
    "workpaper_ref": "調書参照番号",
    "prepared_by": "作成者",
    "reviewed_by": "レビュー者",
    "date": "作成日",
    "objective": "監査目的",
    "scope": "監査範囲",
    "methodology": "監査手法",
    "procedures_performed": [
        {{
            "procedure": "実施した監査手続",
            "result": "結果",
            "conclusion": "結論"
        }}
    ],
    "findings": [
        {{
            "finding_id": "F001",
            "finding_type": "タイプ（所見/発見事項/例外）",
            "description": "発見事項の説明",
            "criteria": "評価基準",
            "condition": "実際の状況",
            "cause": "原因",
            "effect": "影響",
            "recommendation": "推奨事項",
            "management_response": "経営者の対応（該当する場合）",
            "priority": "優先度（高/中/低）"
        }}
    ],
    "conclusion": "総合結論",
    "attachments": ["添付資料リスト"],
    "follow_up_required": true/false
}}

## インタビュー記録
{transcript}
""",
    }

    def __init__(self, ai_provider=None):
        """Initialize the report generator.

        Args:
            ai_provider: AI provider for text generation (optional)
        """
        self.ai_provider = ai_provider

    async def generate_report(
        self,
        interview_id: UUID,
        transcript: str,
        report_type: ReportType,
        additional_context: dict[str, Any] | None = None,
    ) -> GeneratedReport:
        """Generate a report from interview transcript.

        Args:
            interview_id: The interview ID
            transcript: Interview transcript text
            report_type: Type of report to generate
            additional_context: Additional context for generation

        Returns:
            GeneratedReport with structured content
        """
        # Get the appropriate prompt
        prompt_template = self.PROMPTS.get(report_type)
        if not prompt_template:
            raise ValueError(f"Unsupported report type: {report_type}")

        prompt = prompt_template.format(
            transcript=transcript,
            **(additional_context or {}),
        )

        # Generate using AI provider
        if self.ai_provider:
            from grc_ai import ChatMessage

            messages = [ChatMessage(role="user", content=prompt)]
            response = await self.ai_provider.chat(
                messages=messages,
                response_format={"type": "json_object"},
            )
            content = response.content

            # Parse JSON response
            import json

            try:
                parsed_content = json.loads(content)
            except json.JSONDecodeError:
                logger.error(f"Failed to parse AI response as JSON: {content[:500]}")
                parsed_content = {"raw_content": content, "error": "JSON parse error"}
        else:
            # Mock response for testing
            parsed_content = self._generate_mock_content(report_type)

        return GeneratedReport(
            report_type=report_type,
            title=parsed_content.get("title", f"{report_type.value} Report"),
            generated_at=datetime.utcnow(),
            content=parsed_content,
            metadata={
                "interview_id": str(interview_id),
                "transcript_length": len(transcript),
            },
        )

    def _generate_mock_content(self, report_type: ReportType) -> dict[str, Any]:
        """Generate mock content for testing."""
        if report_type == ReportType.SUMMARY:
            return {
                "title": "インタビュー要約",
                "interviewee_role": "経理部マネージャー",
                "date": datetime.now().isoformat(),
                "duration_minutes": 45,
                "key_topics": ["月次決算プロセス", "内部統制"],
                "summary": "月次決算プロセスの詳細についてヒアリングを実施。",
                "key_findings": [
                    {
                        "topic": "月次決算",
                        "finding": "手作業による照合作業が多い",
                        "significance": "中",
                    }
                ],
                "follow_up_items": ["システム自動化の検討"],
                "quotes": ["毎月20時間程度を照合作業に費やしている"],
            }
        elif report_type == ReportType.PROCESS_DOC:
            return {
                "title": "業務記述書",
                "process_name": "月次決算プロセス",
                "process_owner": "経理部長",
                "department": "経理部",
                "objective": "月次財務諸表の適時・正確な作成",
                "scope": "売上計上から財務諸表作成まで",
                "frequency": "月次",
                "narrative": "毎月5営業日までに前月の財務諸表を作成する。",
                "process_steps": [
                    {
                        "step_number": 1,
                        "description": "売上データの取得",
                        "responsible_party": "経理担当者",
                        "system_used": "基幹システム",
                        "inputs": ["販売データ"],
                        "outputs": ["売上仕訳"],
                        "controls": ["データ照合"],
                    }
                ],
                "key_controls": ["承認ワークフロー", "照合チェック"],
                "systems_involved": ["基幹システム", "会計システム"],
                "risks": ["データ入力ミス", "承認遅延"],
                "improvement_opportunities": ["自動化推進"],
            }
        elif report_type == ReportType.RCM:
            return {
                "title": "リスクコントロールマトリックス",
                "process_name": "月次決算プロセス",
                "assessment_date": datetime.now().isoformat(),
                "prepared_by": "監査チーム",
                "items": [
                    {
                        "risk_id": "R001",
                        "risk_description": "売上計上の誤り",
                        "risk_category": "財務報告",
                        "likelihood": "中",
                        "impact": "高",
                        "inherent_risk": "高",
                        "control_id": "C001",
                        "control_description": "売上データの照合",
                        "control_type": "発見的",
                        "control_frequency": "月次",
                        "control_owner": "経理担当者",
                        "control_effectiveness": "有効",
                        "residual_risk": "低",
                        "gap_identified": "なし",
                        "remediation_action": "なし",
                    }
                ],
                "summary": {
                    "total_risks": 1,
                    "high_risks": 0,
                    "medium_risks": 1,
                    "low_risks": 0,
                    "control_gaps": 0,
                    "key_observations": ["統制は概ね有効に機能"],
                },
            }
        else:  # AUDIT_WORKPAPER
            return {
                "title": "監査調書",
                "audit_project": "月次決算プロセス監査",
                "workpaper_ref": "WP-001",
                "prepared_by": "監査担当者",
                "reviewed_by": "監査マネージャー",
                "date": datetime.now().isoformat(),
                "objective": "月次決算プロセスの有効性評価",
                "scope": "経理部門の月次決算業務",
                "methodology": "インタビュー、ウォークスルー",
                "procedures_performed": [
                    {
                        "procedure": "業務プロセスのヒアリング",
                        "result": "プロセスを理解",
                        "conclusion": "特段の問題なし",
                    }
                ],
                "findings": [],
                "conclusion": "月次決算プロセスは適切に運用されている",
                "attachments": ["インタビュー記録"],
                "follow_up_required": False,
            }

    async def export_report(
        self,
        report: GeneratedReport,
        format: ExportFormat,
    ) -> bytes:
        """Export a report to the specified format.

        Args:
            report: The generated report
            format: Export format

        Returns:
            Exported content as bytes
        """
        if format == ExportFormat.JSON:
            import json

            return json.dumps(report.content, ensure_ascii=False, indent=2).encode("utf-8")

        elif format == ExportFormat.MARKDOWN:
            return self._export_to_markdown(report).encode("utf-8")

        elif format == ExportFormat.WORD:
            return await self._export_to_word(report)

        elif format == ExportFormat.EXCEL:
            return await self._export_to_excel(report)

        elif format == ExportFormat.PDF:
            return await self._export_to_pdf(report)

        else:
            raise ValueError(f"Unsupported export format: {format}")

    def _export_to_markdown(self, report: GeneratedReport) -> str:
        """Export report to Markdown format."""
        content = report.content
        lines = [f"# {content.get('title', 'Report')}", ""]

        # Add metadata
        if "date" in content:
            lines.append(f"**Date:** {content['date']}")
        if "prepared_by" in content:
            lines.append(f"**Prepared by:** {content['prepared_by']}")
        lines.append("")

        # Add main content based on report type
        if report.report_type == ReportType.SUMMARY:
            lines.append("## Summary")
            lines.append(content.get("summary", ""))
            lines.append("")

            if "key_findings" in content:
                lines.append("## Key Findings")
                for finding in content["key_findings"]:
                    lines.append(
                        f"- **{finding['topic']}**: {finding['finding']} "
                        f"(Significance: {finding['significance']})"
                    )
                lines.append("")

        elif report.report_type == ReportType.PROCESS_DOC:
            lines.append("## Process Overview")
            lines.append(f"**Process Name:** {content.get('process_name', '')}")
            lines.append(f"**Owner:** {content.get('process_owner', '')}")
            lines.append(f"**Objective:** {content.get('objective', '')}")
            lines.append("")

            if "narrative" in content:
                lines.append("## Process Narrative")
                lines.append(content["narrative"])
                lines.append("")

            if "process_steps" in content:
                lines.append("## Process Steps")
                for step in content["process_steps"]:
                    lines.append(f"### Step {step['step_number']}")
                    lines.append(step["description"])
                    lines.append("")

        elif report.report_type == ReportType.RCM:
            if "items" in content:
                lines.append("## Risk Control Matrix")
                lines.append("")
                lines.append("| Risk ID | Risk Description | Control | Residual Risk |")
                lines.append("|---------|------------------|---------|---------------|")
                for item in content["items"]:
                    lines.append(
                        f"| {item['risk_id']} | {item['risk_description'][:50]} | "
                        f"{item['control_description'][:30]} | {item['residual_risk']} |"
                    )
                lines.append("")

        elif report.report_type == ReportType.AUDIT_WORKPAPER:
            lines.append(f"**Objective:** {content.get('objective', '')}")
            lines.append(f"**Scope:** {content.get('scope', '')}")
            lines.append("")

            if "findings" in content and content["findings"]:
                lines.append("## Findings")
                for finding in content["findings"]:
                    lines.append(f"### {finding['finding_id']}: {finding['description']}")
                    lines.append(f"**Type:** {finding['finding_type']}")
                    if finding.get("recommendation"):
                        lines.append(f"**Recommendation:** {finding['recommendation']}")
                    lines.append("")

            lines.append("## Conclusion")
            lines.append(content.get("conclusion", ""))

        return "\n".join(lines)

    async def _export_to_word(self, report: GeneratedReport) -> bytes:
        """Export report to Word (.docx) format."""
        from io import BytesIO

        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # Style
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Calibri"

        content = report.content
        title = content.get("title", "Report")

        doc.add_heading(title, level=0)

        # Metadata
        meta_fields = {
            "date": "Date",
            "prepared_by": "Prepared by",
            "process_owner": "Process Owner",
            "department": "Department",
        }
        for key, label in meta_fields.items():
            if key in content:
                doc.add_paragraph(f"{label}: {content[key]}")

        if report.report_type == ReportType.SUMMARY:
            if "summary" in content:
                doc.add_heading("Summary", level=1)
                doc.add_paragraph(content["summary"])
            if "key_findings" in content:
                doc.add_heading("Key Findings", level=1)
                for f in content["key_findings"]:
                    doc.add_paragraph(
                        f"{f.get('topic', '')}: {f.get('finding', '')} "
                        f"({f.get('significance', '')})",
                        style="List Bullet",
                    )
            if "follow_up_items" in content:
                doc.add_heading("Follow-up Items", level=1)
                for item in content["follow_up_items"]:
                    doc.add_paragraph(item, style="List Bullet")

        elif report.report_type == ReportType.PROCESS_DOC:
            for field_key, heading in [
                ("objective", "Objective"),
                ("scope", "Scope"),
                ("narrative", "Process Narrative"),
            ]:
                if field_key in content:
                    doc.add_heading(heading, level=1)
                    doc.add_paragraph(str(content[field_key]))
            if "process_steps" in content:
                doc.add_heading("Process Steps", level=1)
                for step in content["process_steps"]:
                    doc.add_heading(
                        f"Step {step.get('step_number', '')}: {step.get('description', '')}",
                        level=2,
                    )
                    if step.get("responsible_party"):
                        doc.add_paragraph(f"Responsible: {step['responsible_party']}")
                    if step.get("system_used"):
                        doc.add_paragraph(f"System: {step['system_used']}")

        elif report.report_type == ReportType.RCM:
            if "items" in content and content["items"]:
                doc.add_heading("Risk Control Matrix", level=1)
                headers = ["Risk ID", "Description", "Control", "Residual Risk"]
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                for item in content["items"]:
                    row = table.add_row()
                    row.cells[0].text = str(item.get("risk_id", ""))
                    row.cells[1].text = str(item.get("risk_description", ""))
                    row.cells[2].text = str(item.get("control_description", ""))
                    row.cells[3].text = str(item.get("residual_risk", ""))

        elif report.report_type == ReportType.AUDIT_WORKPAPER:
            for field_key, heading in [
                ("objective", "Audit Objective"),
                ("scope", "Scope"),
                ("methodology", "Methodology"),
            ]:
                if field_key in content:
                    doc.add_heading(heading, level=1)
                    doc.add_paragraph(str(content[field_key]))
            if "findings" in content and content["findings"]:
                doc.add_heading("Findings", level=1)
                for f in content["findings"]:
                    doc.add_heading(
                        f"{f.get('finding_id', '')}: {f.get('description', '')}",
                        level=2,
                    )
                    if f.get("recommendation"):
                        doc.add_paragraph(f"Recommendation: {f['recommendation']}")
            if "conclusion" in content:
                doc.add_heading("Conclusion", level=1)
                doc.add_paragraph(content["conclusion"])

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def _export_to_excel(self, report: GeneratedReport) -> bytes:
        """Export report to Excel (.xlsx) format."""
        from io import BytesIO

        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill

        wb = Workbook()
        ws = wb.active
        content = report.content

        header_font = Font(bold=True, size=12)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font_white = Font(bold=True, size=11, color="FFFFFF")

        ws.title = content.get("title", "Report")[:31]  # Excel max 31 chars

        if report.report_type == ReportType.RCM and "items" in content:
            columns = [
                ("risk_id", "Risk ID", 12),
                ("risk_description", "Risk Description", 40),
                ("risk_category", "Category", 20),
                ("likelihood", "Likelihood", 12),
                ("impact", "Impact", 12),
                ("control_id", "Control ID", 12),
                ("control_description", "Control Description", 40),
                ("control_type", "Control Type", 15),
                ("control_frequency", "Frequency", 12),
                ("residual_risk", "Residual Risk", 15),
            ]
            for col_idx, (_, header, width) in enumerate(columns, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.font = header_font_white
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
                ws.column_dimensions[chr(64 + col_idx)].width = width
            for row_idx, item in enumerate(content["items"], 2):
                for col_idx, (key, _, _) in enumerate(columns, 1):
                    ws.cell(row=row_idx, column=col_idx, value=str(item.get(key, "")))

        elif report.report_type == ReportType.PROCESS_DOC and "process_steps" in content:
            columns = [
                ("step_number", "Step #", 8),
                ("description", "Description", 50),
                ("responsible_party", "Responsible", 20),
                ("system_used", "System", 20),
            ]
            for col_idx, (_, header, width) in enumerate(columns, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.font = header_font_white
                cell.fill = header_fill
                ws.column_dimensions[chr(64 + col_idx)].width = width
            for row_idx, step in enumerate(content["process_steps"], 2):
                for col_idx, (key, _, _) in enumerate(columns, 1):
                    ws.cell(row=row_idx, column=col_idx, value=str(step.get(key, "")))

        else:
            # Generic key-value export for other report types
            ws.cell(row=1, column=1, value="Key").font = header_font
            ws.cell(row=1, column=2, value="Value").font = header_font
            ws.column_dimensions["A"].width = 25
            ws.column_dimensions["B"].width = 60
            row = 2
            for key, value in content.items():
                ws.cell(row=row, column=1, value=str(key))
                if isinstance(value, list | dict):
                    import json

                    ws.cell(row=row, column=2, value=json.dumps(value, ensure_ascii=False))
                else:
                    ws.cell(row=row, column=2, value=str(value))
                row += 1

        buf = BytesIO()
        wb.save(buf)
        return buf.getvalue()

    async def _export_to_pdf(self, report: GeneratedReport) -> bytes:
        """Export report to PDF format."""
        from io import BytesIO

        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20 * mm, bottomMargin=20 * mm)
        styles = getSampleStyleSheet()
        content = report.content
        elements: list = []

        # Title
        title_style = ParagraphStyle(
            "CustomTitle", parent=styles["Title"], fontSize=18, spaceAfter=12
        )
        elements.append(Paragraph(content.get("title", "Report"), title_style))
        elements.append(Spacer(1, 6 * mm))

        # Metadata
        meta_style = styles["Normal"]
        for key, label in [
            ("date", "Date"),
            ("prepared_by", "Prepared by"),
            ("process_owner", "Process Owner"),
        ]:
            if key in content:
                elements.append(Paragraph(f"<b>{label}:</b> {content[key]}", meta_style))

        elements.append(Spacer(1, 4 * mm))

        heading_style = ParagraphStyle(
            "CustomHeading", parent=styles["Heading2"], fontSize=14, spaceAfter=6
        )

        if report.report_type == ReportType.RCM and "items" in content:
            elements.append(Paragraph("Risk Control Matrix", heading_style))
            table_data = [["Risk ID", "Description", "Control", "Residual"]]
            for item in content["items"]:
                table_data.append([
                    str(item.get("risk_id", "")),
                    str(item.get("risk_description", ""))[:60],
                    str(item.get("control_description", ""))[:40],
                    str(item.get("residual_risk", "")),
                ])
            t = Table(table_data, colWidths=[50, 180, 150, 70])
            t.setStyle(
                TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
                ])
            )
            elements.append(t)
        else:
            # Render markdown-like content for other types
            md = self._export_to_markdown(report)
            for line in md.split("\n"):
                line = line.strip()
                if not line:
                    elements.append(Spacer(1, 2 * mm))
                elif line.startswith("## "):
                    elements.append(Paragraph(line[3:], heading_style))
                elif line.startswith("# "):
                    pass  # skip duplicate title
                elif line.startswith("- ") or line.startswith("* "):
                    bullet_text = line[2:]
                    elements.append(
                        Paragraph(f"\u2022 {bullet_text}", styles["Normal"])
                    )
                elif line.startswith("**") and line.endswith("**"):
                    elements.append(
                        Paragraph(f"<b>{line.strip('*')}</b>", styles["Normal"])
                    )
                else:
                    elements.append(Paragraph(line, styles["Normal"]))

        doc.build(elements)
        return buf.getvalue()


# Singleton instance
_report_generator: ReportGeneratorService | None = None


def get_report_generator() -> ReportGeneratorService:
    """Get or create the report generator singleton."""
    global _report_generator
    if _report_generator is None:
        from grc_backend.api.deps import get_ai_provider
        from grc_backend.config import get_settings

        settings = get_settings()
        try:
            ai_provider = get_ai_provider(settings)
        except Exception:
            ai_provider = None
        _report_generator = ReportGeneratorService(ai_provider=ai_provider)
    return _report_generator
